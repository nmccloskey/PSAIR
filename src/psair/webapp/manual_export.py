from __future__ import annotations

from dataclasses import asdict, dataclass
from html import escape
from importlib.util import find_spec
from io import BytesIO
from pathlib import Path
import re
import shutil
import tempfile
from typing import Mapping

from ..manual.index import ManualFile, build_manual_index, numeric_sort_key
from ..manual.pdf import (
    add_pagebreaks_between_sections,
    build_pandoc_extra_args,
    run_pandoc,
    strip_leading_heading_numbers,
)


class ManualExportError(RuntimeError):
    """Raised when a manual export backend cannot produce a document."""


@dataclass(frozen=True)
class ManualSectionMeta:
    rel_path: str
    title: str
    filename: str


def detect_manual_export_backends(*, check_pdf: bool = True) -> dict[str, bool]:
    """
    Return export backends currently available at runtime.

    Pandoc is an external executable, not a Python dependency. Local CLI users
    need Pandoc installed separately for Pandoc PDF output; hosted Streamlit
    deployments can provide it through the root-level packages.txt. If Pandoc
    is absent, PDF export can fall back to WeasyPrint. DOCX remains independent.
    """
    return {
        "pandoc_pdf": check_pdf and shutil.which("pandoc") is not None,
        "weasyprint_pdf": check_pdf and _weasyprint_pdf_available(),
        "docx": _module_available("docx"),
    }


def get_best_pdf_backend() -> str | None:
    """Return the preferred available PDF backend name."""
    backends = detect_manual_export_backends()
    if backends["pandoc_pdf"]:
        return "pandoc"
    if backends["weasyprint_pdf"]:
        return "weasyprint"
    return None


def build_manual_markdown(
    manual_dir: str | Path,
    *,
    strip_heading_numbers: bool = True,
    pagebreaks: bool = True,
) -> tuple[str, list[dict[str, str]]]:
    """Build export-ready Markdown using the same index ordering as the viewer."""
    manual_root = Path(manual_dir).resolve()
    _tree, flat = build_manual_index(manual_root)
    return build_manual_markdown_from_index(
        flat,
        strip_heading_numbers=strip_heading_numbers,
        pagebreaks=pagebreaks,
    )


def build_manual_markdown_from_index(
    flat: Mapping[str, ManualFile],
    *,
    strip_heading_numbers: bool = True,
    pagebreaks: bool = True,
) -> tuple[str, list[dict[str, str]]]:
    """Concatenate indexed manual sections into one Markdown document."""
    chunks: list[str] = []
    metadata: list[ManualSectionMeta] = []

    for mf in _ordered_manual_files(flat):
        text = mf.text
        if strip_heading_numbers:
            text = strip_leading_heading_numbers(text)

        chunk = text.strip()
        if chunk:
            chunks.append(chunk)

        metadata.append(
            ManualSectionMeta(
                rel_path=mf.rel_path.as_posix(),
                title=mf.title,
                filename=mf.rel_path.name,
            )
        )

    markdown_text = (
        add_pagebreaks_between_sections(chunks)
        if pagebreaks
        else "\n\n".join(chunks).strip() + ("\n" if chunks else "")
    )
    return markdown_text, [asdict(item) for item in metadata]


def export_manual_pdf_pandoc(
    markdown_text: str,
    *,
    title: str,
    yaml_path: str | Path | None = None,
    pdf_engine: str = "xelatex",
    margin: str | None = "1in",
    toc: bool = True,
    toc_depth: int | None = 3,
) -> bytes:
    """Render Markdown to PDF bytes with Pandoc, if Pandoc is on PATH."""
    if shutil.which("pandoc") is None:
        raise ManualExportError("Pandoc is not available on PATH.")

    yaml_file = _existing_optional_path(yaml_path)
    with tempfile.TemporaryDirectory(prefix="psair_manual_export_") as tmpdir:
        tmp_root = Path(tmpdir)
        markdown_path = tmp_root / "manual.md"
        output_path = tmp_root / "manual.pdf"
        markdown_path.write_text(markdown_text, encoding="utf-8")

        try:
            run_pandoc(
                markdown_path=markdown_path,
                output_path=output_path,
                pdf_engine=pdf_engine,
                yaml_path=yaml_file,
                extra_args=_pandoc_export_args(
                    title=title,
                    margin=margin,
                    toc=toc,
                    toc_depth=toc_depth,
                ),
            )
        except Exception as exc:
            raise ManualExportError(f"Pandoc PDF export failed: {exc}") from exc

        return output_path.read_bytes()


def export_manual_pdf_weasyprint(
    markdown_text: str,
    *,
    title: str,
    base_url: str | Path | None = None,
) -> bytes:
    """Render Markdown to PDF bytes with Python-Markdown and WeasyPrint."""
    try:
        import markdown as markdown_lib
        from weasyprint import HTML
    except Exception as exc:
        raise ManualExportError(
            "WeasyPrint PDF export requires the markdown and weasyprint packages."
        ) from exc

    html_body = markdown_lib.markdown(
        _markdown_for_html(markdown_text),
        extensions=["extra", "sane_lists", "codehilite"],
        output_format="html5",
    )
    html_doc = _wrap_printable_html(html_body, title=title)

    try:
        pdf_bytes = HTML(
            string=html_doc,
            base_url=str(base_url) if base_url is not None else None,
        ).write_pdf()
    except Exception as exc:
        raise ManualExportError(f"WeasyPrint PDF export failed: {exc}") from exc

    return bytes(pdf_bytes)


def export_manual_docx(markdown_text: str, *, title: str) -> bytes:
    """Render a conservative Markdown subset to DOCX bytes."""
    try:
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from docx.shared import Pt
    except Exception as exc:
        raise ManualExportError("DOCX export requires the python-docx package.") from exc

    document = Document()
    document.add_heading(title, level=0)
    _configure_docx_styles(document, Pt, WD_STYLE_TYPE)
    _write_docx_markdown(document, markdown_text)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def export_manual_pdf(
    markdown_text: str,
    *,
    title: str,
    yaml_path: str | Path | None = None,
    base_url: str | Path | None = None,
) -> tuple[bytes, str]:
    """Render PDF bytes with Pandoc first, then WeasyPrint as fallback."""
    errors: list[str] = []
    if shutil.which("pandoc") is not None:
        try:
            return (
                export_manual_pdf_pandoc(
                    markdown_text,
                    title=title,
                    yaml_path=yaml_path,
                ),
                "pandoc",
            )
        except Exception as exc:
            errors.append(str(exc))

    if _module_available("weasyprint") and _module_available("markdown"):
        try:
            return (
                export_manual_pdf_weasyprint(
                    markdown_text,
                    title=title,
                    base_url=base_url,
                ),
                "weasyprint",
            )
        except Exception as exc:
            errors.append(str(exc))

    details = " ".join(errors) if errors else "No PDF backend is available."
    raise ManualExportError(f"Could not generate PDF. {details}")


def _module_available(name: str) -> bool:
    return find_spec(name) is not None


def _weasyprint_pdf_available() -> bool:
    if not _module_available("markdown") or not _module_available("weasyprint"):
        return False
    try:
        import markdown  # noqa: F401
        from weasyprint import HTML  # noqa: F401
    except Exception:
        return False
    return True


def _ordered_manual_files(flat: Mapping[str, ManualFile]) -> list[ManualFile]:
    return [
        flat[rel_str]
        for rel_str in sorted(
            flat,
            key=lambda rel: [numeric_sort_key(part) for part in Path(rel).parts],
        )
    ]


def _existing_optional_path(pathlike: str | Path | None) -> Path | None:
    if pathlike is None:
        return None
    path = Path(pathlike).resolve()
    return path if path.exists() else None


def _pandoc_export_args(
    *,
    title: str,
    margin: str | None,
    toc: bool,
    toc_depth: int | None,
) -> list[str]:
    args = build_pandoc_extra_args(margin=margin, toc=toc, toc_depth=toc_depth)
    if title:
        args.extend(["--metadata", f"title={title}"])
    return args


def _markdown_for_html(markdown_text: str) -> str:
    return re.sub(
        r"\n\s*\\newpage\s*\n",
        '\n\n<div class="page-break"></div>\n\n',
        markdown_text,
    )


def _wrap_printable_html(body: str, *, title: str) -> str:
    return f"""<!doctype html>
<html>
<head>
  <meta charset="utf-8">
  <title>{escape(title)}</title>
  <style>
    @page {{ margin: 0.8in; }}
    body {{
      color: #1f2933;
      font-family: Arial, Helvetica, sans-serif;
      font-size: 11pt;
      line-height: 1.45;
    }}
    h1, h2, h3, h4 {{ color: #111827; line-height: 1.2; margin-top: 1.4em; }}
    h1 {{ page-break-before: always; border-bottom: 1px solid #d1d5db; padding-bottom: 0.15in; }}
    h1:first-of-type {{ page-break-before: auto; }}
    pre {{
      background: #f3f4f6;
      border: 1px solid #d1d5db;
      font-family: Consolas, "Courier New", monospace;
      font-size: 9pt;
      overflow-wrap: anywhere;
      padding: 0.12in;
      white-space: pre-wrap;
    }}
    code {{ font-family: Consolas, "Courier New", monospace; font-size: 0.92em; }}
    blockquote {{ border-left: 3px solid #9ca3af; color: #4b5563; margin-left: 0; padding-left: 0.18in; }}
    table {{ border-collapse: collapse; width: 100%; }}
    th, td {{ border: 1px solid #d1d5db; padding: 0.06in; vertical-align: top; }}
    .page-break {{ page-break-before: always; }}
  </style>
</head>
<body>
{body}
</body>
</html>"""


def _configure_docx_styles(document, pt_factory, style_type) -> None:
    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = pt_factory(10.5)

    try:
        code_style = styles["PSAIR Code"]
    except KeyError:
        code_style = styles.add_style("PSAIR Code", style_type.PARAGRAPH)
    code_style.font.name = "Courier New"
    code_style.font.size = pt_factory(9)


def _write_docx_markdown(document, markdown_text: str) -> None:
    in_code = False
    code_lines: list[str] = []
    paragraph_lines: list[str] = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.rstrip()
        if line.strip().startswith(("```", "~~~")):
            if in_code:
                _add_code_block(document, code_lines)
                code_lines = []
            else:
                _flush_docx_paragraph(document, paragraph_lines)
                paragraph_lines = []
            in_code = not in_code
            continue

        if in_code:
            code_lines.append(raw_line)
            continue

        paragraph_lines = _write_docx_markdown_line(document, line, paragraph_lines)

    _add_code_block(document, code_lines)
    _flush_docx_paragraph(document, paragraph_lines)


def _write_docx_markdown_line(document, line: str, pending: list[str]) -> list[str]:
    stripped = line.strip()
    if not stripped:
        _flush_docx_paragraph(document, pending)
        return []
    if stripped == r"\newpage":
        _flush_docx_paragraph(document, pending)
        document.add_page_break()
        return []
    if re.fullmatch(r"[-*_]{3,}", stripped):
        _flush_docx_paragraph(document, pending)
        document.add_paragraph("_" * 28)
        return []

    if _is_structured_markdown_line(stripped):
        _flush_docx_paragraph(document, pending)
        _write_structured_docx_line(document, stripped)
        return []

    return [*pending, _clean_inline_markdown(stripped)]


def _is_structured_markdown_line(stripped: str) -> bool:
    patterns = [
        r"^#{1,6}\s+.+$",
        r"^[-*+]\s+.+$",
        r"^\d+[.)]\s+.+$",
        r"^>\s?.+$",
    ]
    return any(re.match(pattern, stripped) for pattern in patterns)


def _write_structured_docx_line(document, stripped: str) -> None:
    heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
    if heading:
        level = min(len(heading.group(1)), 4)
        document.add_heading(_clean_inline_markdown(heading.group(2)), level=level)
        return

    bullet = re.match(r"^[-*+]\s+(.+)$", stripped)
    if bullet:
        document.add_paragraph(
            _clean_inline_markdown(bullet.group(1)),
            style="List Bullet",
        )
        return

    numbered = re.match(r"^\d+[.)]\s+(.+)$", stripped)
    if numbered:
        document.add_paragraph(
            _clean_inline_markdown(numbered.group(1)),
            style="List Number",
        )
        return

    quote = re.match(r"^>\s?(.+)$", stripped)
    if quote:
        document.add_paragraph(
            _clean_inline_markdown(quote.group(1)),
            style="Intense Quote",
        )


def _flush_docx_paragraph(document, lines: list[str]) -> None:
    text = " ".join(part.strip() for part in lines if part.strip())
    if text:
        document.add_paragraph(text)


def _add_code_block(document, lines: list[str]) -> None:
    if not lines:
        return
    paragraph = document.add_paragraph(style="PSAIR Code")
    for index, line in enumerate(lines):
        paragraph.add_run(line)
        if index < len(lines) - 1:
            paragraph.add_run().add_break()


def _clean_inline_markdown(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.*?)\1", r"\2", text)
    return text
